# order_bot/llm_file_parser.py
from __future__ import annotations

import io
from typing import Any

import openpyxl

try:
    import docx as _docx
except ImportError:
    _docx = None

from order_bot.llm.client import JSONLLMClient
from order_bot.parsers.models import ParseError, ParseResult


# ───────────────────────── конвертеры файл → текст ─────────────────────────

def _xlsx_to_text(file_bytes: bytes) -> str:
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    lines = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() if c is not None else "" for c in row]
            if any(c for c in cells):
                lines.append("\t".join(cells))
    return "\n".join(lines)


def _docx_to_text(file_bytes: bytes) -> str:
    if _docx is None:
        raise RuntimeError("python-docx не встановлено")
    doc = _docx.Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)


def file_to_text(file_bytes: bytes, filename: str) -> str:
    name = filename.lower()
    if name.endswith((".xlsx", ".xls")):
        return _xlsx_to_text(file_bytes)
    if name.endswith(".docx"):
        return _docx_to_text(file_bytes)
    # CSV / plain text
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


# ───────────────────────── промпти ─────────────────────────

_PRICE_SYSTEM = """Ти — парсер прайс-листів агрохімікатів. Рядки розділені табуляцією.

  Формат: № з/п | Назва препарату | Базова ціна | Ціна при 200 тис | Знижка 200 тис,% | Ціна при 150 тис | Знижка 150 тис,% | Ціна при 100 тис | Знижка 100 тис,%

  Витягни для кожного:
  - name: Назва препарату з фасуванням (напр. "АВАТАР (каністра 20 л)")
  - base_price: Базова ціна (2-га числова колонка)
  - price_200k: Ціна при 200 тис (3-я числова колонка)
  - discount_200k_percent: Знижка при 200 тис (4-та колонка, число без %)
  - price_150k: Ціна при 150 тис (5-та числова колонка)
  - discount_150k_percent: Знижка при 150 тис (6-та колонка)
  - price_100k: Ціна при 100 тис (7-ма числова колонка)
  - discount_100k_percent: Знижка при 100 тис (8-ма колонка)

  Приклад: "1 АВАТАР (каністра 20 л) 3.95 3.85 2.53% 3.87 2.03% 3.91 1.01%"

  ПОВЕРНИ ТІЛЬКИ ВАЛІДНИЙ JSON:
  {"rows": [{"name": "...", "base_price": X.XX, "price_200k": X.XX, "discount_200k_percent": X.XX, "price_150k": X.XX, "discount_150k_percent": X.XX, "price_100k": X.XX, "discount_100k_percent": X.XX}],
  "errors": []}""".strip()

_STOCK_SYSTEM = """
Ти — парсер залишків товарів.
Знайди всі позиції: назва товару + кількість (qty).
Поверни ТІЛЬКИ валідний JSON:
{"rows": [{"name": "...", "qty": 10}], "errors": []}
""".strip()

_WAREHOUSE_SYSTEM = """
Ти — парсер довідника складів/локацій.
Знайди всі склади: назва (і адреса якщо є).
Поверни ТІЛЬКИ валідний JSON:
{"rows": [{"name": "...", "address": "..."}], "errors": []}
""".strip()

_SYSTEMS: dict[str, str] = {
    "price": _PRICE_SYSTEM,
    "stock": _STOCK_SYSTEM,
    "warehouse": _WAREHOUSE_SYSTEM,
}

# ───────────────────────── нормалізація рядків ─────────────────────────

def _normalize_price_row(raw: dict, idx: int) -> tuple[dict[str, Any] | None, ParseError | None]:
    name = str(raw.get("name") or "").strip()
    if not name:
        return None, ParseError(row=idx, field="name", message="Відсутня назва препарату")

    sku = str(raw.get("sku") or name).strip()

    try:
        base_price = float(raw.get("base_price") or 0)
    except (ValueError, TypeError):
        return None, ParseError(row=idx, field="base_price", message=f"Невалідна базова ціна: {raw.get('base_price')}")

    if base_price <= 0:
        return None, ParseError(row=idx, field="base_price", message="Базова ціна має бути > 0")

    def _float(val):
        try:
            return float(val or 0)
        except (ValueError, TypeError):
            return 0.0

    return {
        "sku": sku,
        "name": name,
        "base_price": base_price,
        "price_200k": _float(raw.get("price_200k")),
        "discount_200k_percent": _float(raw.get("discount_200k_percent")),
        "price_150k": _float(raw.get("price_150k")),
        "discount_150k_percent": _float(raw.get("discount_150k_percent")),
        "price_100k": _float(raw.get("price_100k")),
        "discount_100k_percent": _float(raw.get("discount_100k_percent")),
    }, None


def _normalize_stock_row(raw: dict, idx: int) -> tuple[dict[str, Any] | None, ParseError | None]:
    name = str(raw.get("name") or "").strip()
    sku = str(raw.get("sku") or name or "").strip()
    if not sku:
        return None, ParseError(row=idx, field="sku", message="Відсутня назва товару/артикул")

    qty_val = raw.get("qty") or raw.get("quantity")
    try:
        quantity = int(float(qty_val or 0))
    except (ValueError, TypeError):
        return None, ParseError(row=idx, field="quantity", message=f"Невалідна кількість: {qty_val}")

    return {"sku": sku, "quantity": quantity}, None


def _normalize_warehouse_row(raw: dict, idx: int) -> tuple[dict[str, Any] | None, ParseError | None]:
    name = str(raw.get("name") or "").strip()
    if not name:
        return None, ParseError(row=idx, field="name", message="Відсутня назва складу")

    return {
        "name": name,
        "address": str(raw.get("address") or "").strip() or None,
    }, None


_NORMALIZERS: dict[str, Any] = {
    "price": _normalize_price_row,
    "stock": _normalize_stock_row,
    "warehouse": _normalize_warehouse_row,
}


# ───────────────────────── основний парсер ─────────────────────────

class LLMFileParser:
    MAX_CHARS = 120_000  # ~30k токенів, з запасом

    def __init__(self, llm: JSONLLMClient | None = None) -> None:
        self._llm = llm

    @staticmethod
    def _detect_type_by_name(filename: str) -> str:
        text = filename.lower()
        if any(token in text for token in ["price", "pricelist", "прайс"]):
            return "price"
        if any(token in text for token in ["warehouse", "warehouses", "локац", "branch", "склад"]):
            return "warehouse"
        if any(token in text for token in ["stock", "остат", "залишк", "quantity"]):
            return "stock"
        return "price"

    def parse(
        self,
        file_bytes: bytes,
        filename: str,
        forced_type: str | None = None,  # "price" | "stock" | "warehouse"
    ) -> ParseResult:
        doc_type = forced_type or self._detect_type_by_name(filename)
        system_prompt = _SYSTEMS.get(doc_type, _PRICE_SYSTEM)

        raw_text = file_to_text(file_bytes, filename)
        if len(raw_text) > self.MAX_CHARS:
            raw_text = raw_text[: self.MAX_CHARS] + "\n...[truncated]"

        if self._llm is None or not self._llm.enabled:
            return ParseResult(
                rows=[],
                errors=[ParseError(row=0, field="llm", message="LLM client is not configured")],
            )

        try:
            data = self._llm.parse_json(system_prompt, raw_text)
            print(f"🔍 LLM RAW ROWS (перші 3): {data.get('rows', [])[:3]}", flush=True)
        except Exception as exc:
            return ParseResult(rows=[], errors=[ParseError(row=0, field="llm", message=f"LLM error: {exc}")])

        llm_errors: list[str] = data.get("errors") or []
        raw_rows: list[dict] = data.get("rows") or []

        rows: list[dict[str, Any]] = []
        errors: list[ParseError] = []

        normalizer = _NORMALIZERS.get(doc_type, _normalize_price_row)

        for idx, r in enumerate(raw_rows):
            normalized_row, err = normalizer(r, idx)
            if err is not None:
                errors.append(err)
                continue
            if normalized_row is not None:
                rows.append(normalized_row)

        for msg in llm_errors:
            errors.append(ParseError(row=0, field="llm", message=msg))

        return ParseResult(rows=rows, errors=errors)
